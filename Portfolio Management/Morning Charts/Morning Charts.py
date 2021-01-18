
# Libraries

import numpy as np
import pandas as pd
import pandas_datareader as dr
import matplotlib.pyplot as plt
from docx import Document
from datetime import date
from docx.shared import Inches
from docx.shared import RGBColor
from scipy.stats import ttest_ind
from docx.enum.text import WD_ALIGN_PARAGRAPH
global p

# Functions

def slope(dfp):

    ' Get data'
    datosp = pd.DataFrame(data={})
    datosp = dfp.copy()
    y = datosp.iloc[0:].to_numpy()
    x = np.arange(len(y))
    
    n = len(y)
    sumax = sum(x)
    sumay = sum(y)
    sumax2 = sum(np.power(x,2))
    sumaxy = sum(x*y)

    'Calculate slope'
    m = (sumax*sumay - n*sumaxy)/(sumax**2 - n*sumax2)

    return m   

def SupportResistances(df):
    
    'Define variables'
    df['Periodo']=np.nan
    df['PP'] = np.nan
    df['R1'] = np.nan
    df['S1'] =np.nan
    
    'Initialize lists'
    pivots =[]
    dates = []
    counter = 0
    lastPivot = 0
    Range = [0,0,0,0,0,0,0,0,0,0]
    daterange = [0,0,0,0,0,0,0,0,0,0]
    inicio=0
    
    'This loop calculates the supports'
    for i in df.index:
        
        currentMin = min(Range , default=100000000)
        value=df["Adj Close"][i]
        Range=Range[1:9]
        Range.append(value)
        daterange=daterange[1:9]
        daterange.append(df['Date'][i])
        
        if currentMin == min(Range , default=100000000):
            counter+=1
            ind=i
        else:
            counter = 0
            df.loc[i,('Periodo')]=np.nan
        if counter == 5:
            if currentMin!=0:
                df.loc[inicio:i,('PP')] = (df.loc[inicio:i,('High')] + df.loc[inicio:i,('Low')]+ df.loc[inicio:i,('Adj Close')])/ 3
                df.loc[inicio:i,('S1')] = 2*df.loc[inicio:i,('PP')]-df.loc[inicio:i,('High')]
                df.loc[i,('Periodo')]=currentMin
                lastPivot=currentMin
                dateloc =Range.index(lastPivot)
                lastDate = daterange[dateloc]
                pivots.append(lastPivot)
                dates.append(lastDate)
                if min(df.loc[inicio:i,('S1')])<( min(df.loc[inicio:i,('Low')])   ):
                    df.loc[inicio:i-1,('S1')] = min(df.loc[inicio:i,('Adj Close')])
                    df.loc[i-1:i,('S1')]=np.nan
                else:
                    df.loc[inicio:i-1,('S1')] = min(df.loc[inicio:i,('Adj Close')])
                    df.loc[i-1:i,('S1')]=np.nan
                inicio=ind+1
    
    'Initialize lists'
    pivots =[]
    dates = []
    counter = 0
    lastPivot = 0
    Range = [0,0,0,0,0,0,0,0,0,0]
    daterange = [0,0,0,0,0,0,0,0,0,0]
    inicio=0
    
    'This loop calculates the resistances'
    for i in df.index:
        currentMax = max(Range , default=0)
        value=df["Adj Close"][i]
        Range=Range[1:9]
        Range.append(value)
        daterange=daterange[1:9]
        daterange.append(df['Date'][i])
        if currentMax == max(Range , default=0):
            counter+=1
            ind=i
        else:
            counter = 0
            df.loc[i,('Periodo')]=np.nan
        if counter == 5:
            if currentMax!=0:
                df.loc[inicio:i,('PP')]= (df.loc[inicio:i,('High')] + df.loc[inicio:i,('Low')]+ df.loc[inicio:i,('Adj Close')])/ 3
                df.loc[inicio:i,('R1')] = 2 * df.loc[inicio:i,('PP')] - df.loc[inicio:i,('Low')]
                df.loc[i,('Periodo')]=currentMax
                lastPivot=currentMax
                dateloc =Range.index(lastPivot)
                lastDate = daterange[dateloc]
                pivots.append(lastPivot)
                dates.append(lastDate)
                if df['R1'][inicio:i].max()>(max(df['High'][inicio:i])):
                    df.loc[inicio:i-1,('R1')]=max(df.loc[inicio:i,('Adj Close')])
                    df.loc[i-1:i,('R1')]=np.nan
                else:
                    df.loc[inicio:i-1,('R1')]=max(df.loc[inicio:i,('Adj Close')])
                    df.loc[i-1:i,('R1')]=np.nan
                inicio=ind+1
    return df

def charts(df,stock_ticker,txt1,txt2,best_sma):
    
    difmin=100000000
    valorult=0
    ctend=0
    
    ctend2=slope(df["Adj Close"].tail(21))
    if ctend2>.03:
        ctend=5
        tendencia="Positive trend"
    elif ctend2<-.03:
        ctend=-5
        tendencia="Negative trend"
    else:
        ctend=0
        
    for res in range(len(df["R1"])):
        if ctend<-3 and df["Adj Close"].tail(1).item()>df["S1"][res].item():
            d2=abs(df["Adj Close"].tail(1).item()-df["S1"][res].item())
            if d2<difmin:
                difmin=d2
                colr = '#00CC00'
                valorult=df["S1"][res].item()
                res_or_sop='Support'
                trend_kind = 'Negative'
                color_trend = RGBColor(255, 0, 0)
        elif ctend>3 and df["Adj Close"].tail(1).item()<df["R1"][res].item():
            d1=abs(df["Adj Close"].tail(1).item()-df["R1"][res].item())
            if d1<difmin:
                difmin=d1   
                colr = '#FF0000'
                valorult=df["R1"][res].item()
                res_or_sop = 'Resistance'
                trend_kind = 'Positive'
                color_trend = RGBColor(26,184, 26)
    
    # Build final chart
    
    'Adjusted Close, SMAs, Supposrts and Resistances'
    plt.figure(figsize=(10,18), facecolor='w')
    ax1 = plt.subplot2grid((6, 2), (0, 0), colspan=2,rowspan=2)
    ax1.plot(df['Adj Close'],color='#080808', lw=1.5)
    ax1.plot(df['21SMA'],color='#32649D',lw=1.5)
    ax1.plot(df[str(best_sma)+'SMA'],color='#ED7D31',lw=1.5)
    ax1.plot(df['Bollinup'],color='#404040',lw=1)
    ax1.plot(df['Bollind'],color='#404040',lw=1)
    ax1.plot(df['S1'],color='#00CC00',lw=.9)
    ax1.plot(df['R1'],color='#FF0000',lw=.9)
    if valorult!=0:
        df["sigind"]=valorult
        ax1.plot(df['sigind'],color=colr,lw=.9)
    
    'Bollinger'
    ax1.fill_between(np.arange(len(df['Bollind'])),df['Bollinup'],df['Bollind'],facecolor='grey', alpha=0.2)
    ax1.set_title(stock_ticker)
    ax1.set_ylabel('Adj Close ')
    ax1.get_xaxis().set_ticks([])
    plt.gca().legend(('Adj Close','21SMA',str(best_sma)+'SMA'))

    'Volume'
    ax2 = plt.subplot2grid((6, 2), (2, 0) , colspan=2)
    ax2.bar(range(len(df['Volume'])),df['Volume'],color='#ED7D31')
    ax2.plot(range(len(df['Volume'])),df['Volume'].rolling(window=21).mean(),color='#32649D')
    ax2.set_ylabel('Volume')
    ax2.get_xaxis().set_ticks([])

    'RSI'
    ax3 = plt.subplot2grid((6, 2), (3, 0), colspan=2)
    ax3.plot(df['RSI14'],color='#080808',lw=1.5,label='RSI14')
    ax3.plot(df[txt1],color='#FF0000',lw=1.5,label=txt1)
    ax3.plot(df[txt2],color='#00CC00',lw=1.5,label=txt2)
    ax3.set_ylabel('RSI14')
    ax3.get_xaxis().set_ticks([])
    ax3.legend(loc='upper left')

    'MACD'
    ax4 = plt.subplot2grid((6, 2), (4, 0), colspan=2)
    dates = [ df['Date'][i] for i in range(0,len(df['Date']),5) ]
    exp1 = df['Adj Close'].ewm(span=12, adjust=False).mean()
    exp2 = df['Adj Close'].ewm(span=26, adjust=False).mean()
    macd = exp1-exp2
    exp3 = macd.ewm(span=9, adjust=False).mean()
    diff_macd = macd - exp3
    df['MACD'] = macd
    df['Signal 9'] = exp3
    df['Diff'] = diff_macd
    ax4.plot(df['Date'], macd, color='#080808')
    ax4.plot(df['Date'], exp3, color='#FF0000')
    ax4.bar(df['Date'], diff_macd, color='#00CC00')
    ax4.set_ylabel('MACD level')
    plt.gca().legend(('MACD','Signal Line'))
    
    'Save chart as .jpg'
    plt.savefig('Last Chart.jpg',bbox_inches='tight')
    
    return valorult

def BestSMA(data):
    
    'Period for the return we want to explain'
    days_fwd = 5 
    
    'Percentage for training'
    train_size = 0.60
    
    'Shift the data days_fwd and calculates return'
    data['Forward Adj Close'] = data['Adj Close'].shift(-days_fwd)
    data['Forward Return'] = (data['Forward Adj Close'] - data['Adj Close'])/data['Adj Close']
    
    result = []
    
    'Find best SMA'
    for sma_length in range(21, 127):
        
        data['SMA'] = data['Adj Close'].rolling(sma_length).mean()
        data['input'] = [int(x) for x in data['Adj Close'] > data['SMA']]

        df = data.dropna()
        
        training = df.head(int(train_size * df.shape[0]))
        test = df.tail(int((1 - train_size) * df.shape[0]))

        tr_returns = training[training['input'] == 1]['Forward Return']
        test_returns = test[test['input'] == 1]['Forward Return']
        
        mean_forward_return_training = tr_returns.mean()
        mean_forward_return_test = test_returns.mean()
        
        pvalue = ttest_ind(tr_returns,test_returns,equal_var=False)[1]
        
        result.append({
          'sma_length':sma_length,
          'training_forward_return': mean_forward_return_training,
          'test_forward_return': mean_forward_return_test,
          'p-value':pvalue
        })
    
    'Sort returns'
    result.sort(key = lambda x : -x['training_forward_return'])

    'Return best SMA'
    return result[0]['sma_length']

def rsi(prices, n=14):
    
    'This function calculates the RSI'
    deltas = np.diff(prices)
    seed = deltas[:n]
    up = seed[seed>=0].sum()/n
    down = -seed[seed<0].sum()/n
    rs = up/down
    rsi = np.zeros_like(prices)
    rsi[:n+1] = 100. - 100./(1.+rs)
    
    for i in range(n+1, len(prices)):
        delta = deltas[i-1]
        if delta>0:
            upval = delta
            downval = 0.
        else:
            upval = 0.
            downval = -delta
        up = (up*(n-1) + upval)/n
        down = (down*(n-1) + downval)/n
        rs = up/down
        rsi[i] = 100. - 100./(1.+rs)
    return rsi

# Main

'Read Global Tickers list'
Names = pd.read_excel('Global_Tickers.xlsx').dropna()

'Asign RSI and traded value '
max_rsi = 70
min_rsi = 30
traded_value_mult = 1

'Creates an empty document for charts'
document = Document()

'Add style to the document'
style = document.styles['Normal']
font = style.font
font.name = 'Calibri'

'Adds title to the document'
document.add_heading('Global Morning Charts', 0)

'Creates initial menu'
print("Do you want to run the excel list or just one ticker?")
bandera=int(input("1-Excel list\n2-Single ticker\n="))
if bandera==2:
    stock_ticker=input("Enter ticker\n=")
    if stock_ticker!="":
        stock_name=input("Add name\n=")

if bandera==2:
    try:
        today = date.today()
        prices = dr.DataReader(stock_ticker,'yahoo',str(today.day)+'/'+str(today.month)+'/'+str(today.year-3))
        df = pd.DataFrame(prices)
        df['Date'] = df.index
        df.index = range(0,len(df))
    except:
        print('No data for ',stock_ticker )
    
    'Creates a column with 1D Return'
    df['1D Return'] = np.log(df['Adj Close']/df['Adj Close'].shift(1)).dropna()
    df['1D Return'].fillna(0,inplace=True)
    
    'Calculate RSI'
    df['RSI14'] = rsi(df['Adj Close'])
    
    'Traded Value, here you can select if you want it to be rolling'
    df['Traded Value'] = df['Adj Close']*df['Volume'].rolling(window=2).mean()
    
    'Traded Value SMA21'
    df['Traded Value SMA21'] = df['Traded Value'].rolling(window=21).mean()
    
    '21SMA'
    df['21SMA'] = df['Adj Close'].rolling(window=21).mean()
    best_sma = BestSMA(df )
    df[str(best_sma)+'SMA'] = df['Adj Close'].rolling(window=best_sma).mean()
    
    'RSI parameters'
    df[str(min_rsi)] = 70
    df[str(max_rsi)] = 30
    
    'Bollinger 2 stdev'
    df['Bollinup'] = df['21SMA']+ df['Adj Close'].rolling(window=21).std()*2
    df['Bollind'] = df['21SMA']- df['Adj Close'].rolling(window=21).std()*2
    
    'Supports and resistances'
    df = SupportResistances(df)
    
    'Create chart'
    sigval=charts(df,stock_ticker,str(min_rsi),str(max_rsi), best_sma)

    print('\n')
    print(stock_ticker)
    print(stock_name)
    ctend2=slope(df["Adj Close"].tail(21))
    
    tendencia="No trend"
    if ctend2>.03:
        ctend=5
        tendencia="Positive trend"
    elif ctend2<-.03:
        ctend=-5
        tendencia="Negative trend"
    else:
        ctend=0
    print(tendencia)

    if ctend==5 and sigval!=0:
        print('Next resistance '+str(np.round(sigval,2)))
    elif ctend==-5 and sigval!=0:
        print('Next support '+str(np.round(sigval,2)))                           
    
    'Print Traded value $'
    a_p='2D ADTV $'+str(format(df.loc[len(df)-1,('Traded Value')],',.0f'))+" which is "
    a_p+=str(format(((df.loc[len(df)-1,('Traded Value')]/df.loc[len(df)-1,('Traded Value SMA21')])-1)*100,',.0f'))+'%'
    a_p+=' above its 21D ADTV'
    print(a_p)
    
    'Print RSI' 
    print('RSI14D '+str(format(df.loc[len(df)-1,('RSI14')],',.2f')))

    'Print MACD'
    print('MACD '+str(format(df.loc[len(df)-1,('MACD')],',.2f')+" / Signal(9) "+format(df.loc[len(df)-1,('Signal 9')],',.2f')+" / Diff "+format(df.loc[len(df)-1,('Diff')],',.2f')))
        
else:
    for i in range(len(Names)):
        
        stock_ticker = Names['Tickers'][i]
        stock_name = Names['Name'][i]
        try:
            today = date.today()
            prices = dr.DataReader(stock_ticker,'yahoo',str(today.day)+'/'+str(today.month)+'/'+str(today.year-3))
            df = pd.DataFrame(prices)
            df['Date'] = df.index
            df.index = range(0,len(df))
        except:
            print('No data for ',stock_ticker )
            continue
        
        'Creates a column with 1D Return'
        df['1D Return'] = np.log(df['Adj Close']/df['Adj Close'].shift(1)).dropna()
        df['1D Return'].fillna(0,inplace=True)
        
        'Calculate RSI'
        df['RSI14'] = rsi(df['Adj Close'])
        
        'Traded Value, here you can select if you want it to be rolling'
        df['Traded Value'] = df['Adj Close']*df['Volume'].rolling(window=2).mean()
        
        'Traded Value SMA21'
        df['Traded Value SMA21'] = df['Traded Value'].rolling(window=21).mean()
        
        '21SMA'
        df['21SMA'] = df['Adj Close'].rolling(window=21).mean()
        best_sma = BestSMA(df )
        df[str(best_sma)+'SMA'] = df['Adj Close'].rolling(window=best_sma).mean()
        
        'RSI parameters'
        df[str(min_rsi)] = 70
        df[str(max_rsi)] = 30
        
        'Bollinger 2 stdev'
        df['Bollinup'] = df['21SMA']+ df['Adj Close'].rolling(window=21).std()*2
        df['Bollind'] = df['21SMA']- df['Adj Close'].rolling(window=21).std()*2
        
        'Supports and resistances'
        df = SupportResistances(df)
    
        # Filters
        
        'Traded value'
        if  df.loc[len(df)-1,('Traded Value')] > 10000000:
            
            'RSI'
            if df.loc[len(df)-1,('RSI14')] > max_rsi or df.loc[len(df)-1,('RSI14')] < min_rsi:
    
                'Traded value * traded_value_mult vs SMA21'
                if  df.loc[len(df)-1,('Traded Value')] > df.loc[len(df)-1,('Traded Value SMA21')]*traded_value_mult:

                    p = document.add_paragraph()
                    p.add_run(stock_ticker).bold = True
                    p.add_run('\n')
                    p.add_run(stock_name)
                    
                    # Create charts
                    sigval=charts(df,stock_ticker,str(min_rsi),str(max_rsi), best_sma)
                    ctend2=slope(df["Adj Close"].tail(21))
                    
                    tendencia="No trend"
                    if ctend2>.03:
                        ctend=5
                        tendencia="Positive trend"
                    elif ctend2<-.03:
                        ctend=-5
                        tendencia="Negative trend"
                    else:
                        ctend=0
                    p.add_run('\n'+tendencia)

                    if ctend==5 and sigval!=0:
                        p.add_run('\nNext resistance '+str(np.round(sigval,2)))
                    elif ctend==-5 and sigval!=0:
                        p.add_run('\nNext support '+str(np.round(sigval,2)))                           
                    
                    p.add_run('\n')
                    
                    'Print Traded value $'
                    p.add_run('2D ADTV $')
                    p.add_run(str(format(df.loc[len(df)-1,('Traded Value')],',.0f'))+" which is ")
                    p.add_run(str(format(((df.loc[len(df)-1,('Traded Value')]/df.loc[len(df)-1,('Traded Value SMA21')])-1)*100,',.0f'))+'%').bold = True
                    p.add_run(' above its 21D ADTV\n')
                   
                    'Print RSI' 
                    p.add_run('RSI14D ')
                    p.add_run(str(format(df.loc[len(df)-1,('RSI14')],',.2f')))
    
                    'Print MACD'
                    p.add_run('\nMACD ')
                    p.add_run(str(format(df.loc[len(df)-1,('MACD')],',.2f')+" / Signal(9) "+format(df.loc[len(df)-1,('Signal 9')],',.2f')+" / Diff "+format(df.loc[len(df)-1,('Diff')],',.2f')))
                    
                    'Add images to document in word'
                    my_image = document.add_picture('Last Chart.jpg', width=Inches(4), height=Inches(6))
                    last_paragraph = document.paragraphs[-1]
                    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
         
    # Save final document in word
    document.save('Morning Charts.docx')
        
# Reference

'https://towardsdatascience.com/an-algorithm-to-find-the-best-moving-average-for-stock-trading-1b024672299c'
'https://towardsdatascience.com/detection-of-price-support-and-resistance-levels-in-python-baedc44c34c9'
'https://towardsdatascience.com/pivot-points-calculation-in-python-for-day-trading-659c1e92d323'
