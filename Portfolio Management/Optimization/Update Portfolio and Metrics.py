
# Libraries

import pandas as pd
import numpy as np
import matplotlib.pyplot as plt
from mlfinlab.portfolio_optimization import MeanVarianceOptimisation
import pandas_datareader as dr
from math import sqrt
from sklearn.cluster import KMeans
from matplotlib import pyplot as plt
from numpy.core.shape_base import vstack
from scipy.cluster.hierarchy import dendrogram, linkage
from scipy.cluster.vq import kmeans,vq
import openpyxl
from scipy.stats import kurtosis
import scipy.stats as stats
import seaborn as sns
import mlfinlab as ml
from pathlib import Path
import pathlib
import math
global Tmuestra
import docx
doc=docx.Document()
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Read and update porfolio.xlsx file and download historical data

rutac=pathlib.Path(__file__).parent.absolute()
rutac=str(rutac)
excel_path = Path( rutac+'\porfolio.xlsx')
wb = openpyxl.load_workbook(excel_path)
ws = wb["porfolio"]

def agregaparr(me,doc):
    me=str(me)
    paragraph=doc.add_paragraph(me)
    paragraph.style = doc.styles['Normal']
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    paragraph.paragraph_format.line_spacing = 1
    paragraph.paragraph_format.line_spacing = 1
    doc.save("GraphsResults.docx")
    
def act(columna,stock_prices,wb,ws,rutac):
    ultimop=stock_prices.tail(1)
    penul=stock_prices.tail(2)
    
    columexc=np.arange(30)
    filaexc=np.arange(Tmuestra)
    contador=0
    
    for i in range(len(columexc)):
        if i>0:
            if ws.cell(row=1, column=i).value==columna:
                for j in range(len(filaexc)):
                    if j>0:
                        try:
                            tic= ws.cell(row=j, column=2).value
                            if math.isnan(ultimop[tic].item()):
                                tic= ws.cell(row=j, column=2).value
                                ws.cell(row=j, column=i).value=penul[tic][0].item()
                                wb.save(excel_path)
                            else:             
                                tic= ws.cell(row=j, column=2).value
                                ws.cell(row=j, column=i).value=ultimop[tic].item()
                                wb.save(excel_path)
                        except:
                            print('Print ticker to detect any issue: ', tic)
                    else:
                        pass
        else:
            print('Print row to detect any issue: ', i)
            pass
        
def moments(columna,ultimop,wb,ws,rutac):
    
    columexc=np.arange(30)
    
    filaexc=np.arange(Tmuestra)
    
    contador=0
    
    for i in range(len(columexc)):
        if i>0:
            if ws.cell(row=1, column=i).value==columna:
                for j in range(len(filaexc)):
                    if j>0:
                        try:
                            tic= ws.cell(row=j, column=2).value
                            ws.cell(row=j, column=i).value=ultimop[tic].item()
                            wb.save(excel_path)
                        except:
                            print('Print ticker to detect any issue: ', tic)
                    else:
                        pass
        else:
            print('Print row to detect any issue: ', i)
            pass

'Read tickers'
tickers = pd.read_excel('porfolio.xlsx')['Ticker'].dropna()
tickers.drop(tickers.tail(1).index, inplace = True) 

'Void list'
prices_list = []

'Creates a DF with the data from Yahoo and update the stock prices file'
for ticker in tickers:
    try:
        prices = dr.DataReader(ticker,'yahoo','01/01/2016')['Adj Close']
        prices = pd.DataFrame(prices)
        prices.columns = [ticker]
        prices_list.append(prices)
    except:
        pass
    prices_df = pd.concat(prices_list,axis=1)
prices_df.sort_index(inplace=True)
prices_df.to_excel('stock_prices.xlsx')
stock_prices = prices_df.sort_values(by='Date')

'Sample size'
mt=stock_prices.T
Tmuestra=len(mt)+2

# Update prices and metrcis in porfolio.xlsx file

'Update last price'
act("Price",stock_prices,wb,ws,rutac)


'Generate returns matrix'
ln_matrix = pd.DataFrame(data={})                             
for k in range(len(tickers)):
    ln_matrix[tickers[k]] = np.diff(np.log(prices_df[tickers[k]]))

'Calculate average annual percentage return and volatilities'
etfs_moments = pd.DataFrame(data={}) 
etfs_moments['Mean'] = ln_matrix.mean()
etfs_moments['Variance'] = ln_matrix.var()
etfs_moments['Skew'] = ln_matrix.skew()
etfs_moments['Kurtosis'] = ln_matrix.kurt()

# Function to save them in excel

for etf in etfs_moments.columns:
    moments(etf,etfs_moments[etf],wb,ws,rutac)

for j in tickers:
    print(j)
    print('Mean :', ln_matrix[j].mean())
    print('Variance :', ln_matrix[j].var())
    print('Skew :', ln_matrix[j].skew())
    print('Kurtosis :', ln_matrix[j].kurt())
    h = np.asarray(ln_matrix[j].dropna())
    h = sorted(h)
    
    fit = stats.norm.pdf(h, np.mean(h), np.std(h))
    fig=plt.figure(figsize=(6, 3), facecolor='w')
    plt.plot(h, fit,'-',linewidth = 2)
    plt.hist(h, bins=100)
    plt.show()
    
    # Save variables to send to docx
    
    name=j
    me='Mean :', ln_matrix[j].mean()
    vari='Variance :', ln_matrix[j].var()
    ske='Skew :', ln_matrix[j].skew()
    kurti='Kurtosis :', ln_matrix[j].kurt()
    
    h = np.asarray(ln_matrix[j].dropna())
    h = sorted(h)
    fit = stats.norm.pdf(h, np.mean(h), np.std(h))
    plt.plot(h, fit,'-',linewidth = 2)
    plt.hist(h, bins=100)
    plt.show()
    agregaparr(name,doc)
    agregaparr(me,doc)
    agregaparr(vari,doc)
    agregaparr(ske,doc)
    agregaparr(kurti,doc)
    
    # Call function to add it with the picture
    
    fig.savefig('matplotlibExample.png')
    
    # Use Docx to insert the PNG and save report
    
    my_image = doc.add_picture('matplotlibExample.png')  
    last_paragraph = doc.paragraphs[-1] 
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.save("GraphsResults.docx")

# Clusters of the investable universe

'Calculate average annual percentage return and volatilities over a theoretical one year period'
returns = prices_df.pct_change().mean() * 252
returns = pd.DataFrame(returns)
returns.columns = ['Returns']
returns['Volatility'] = prices_df.pct_change().std() * sqrt(252)

'Format the data as a numpy array to feed into the K-Means algorithm'
data = np.asarray([np.asarray(returns['Returns']),np.asarray(returns['Volatility'])]).T
X = data
distorsions = []
for k in range(2, 15):
    k_means = KMeans(n_clusters=k)
    k_means.fit(X)
    distorsions.append(k_means.inertia_)
fig = plt.figure(figsize=(15, 5))
plt.plot(range(2, 15), distorsions)
plt.grid(True)
plt.title('Elbow curve')
plt.show()

'Computing K-Means with K = 5 (5 clusters)'
centroids,_ = kmeans(data,5)

'Assign each sample to a cluster'
idx,_ = vq(data,centroids)
ultimop=stock_prices.tail(1)
ultimop=ultimop.T
df2 = pd.DataFrame(index=ultimop.index)
df2["clusters"]=idx
df2=df2.T
moments("Cluster",df2,wb,ws,rutac)

fig, ax = plt.subplots(figsize=(10,10))    
for pais in range(len(data)):
    ax.annotate(tickers[pais], (data[pais,0], data[pais,1]), xytext=(0,1),fontsize = 14, textcoords='offset points')
  
'Some plotting using numpys logical indexing'
ax.plot(data[idx==0,0],data[idx==0,1],'ob',
     data[idx==1,0],data[idx==1,1],'oy',
     data[idx==2,0],data[idx==2,1],'or',
     data[idx==3,0],data[idx==3,1],'og',
     data[idx==4,0],data[idx==4,1],'om',)
plt.plot(centroids[:,0],centroids[:,1],'*',color='#ff2e7f',markersize=8)
plt.ylabel('Avg. Volatility Annualized', size=12)
plt.xlabel('Avg. Returns Annualized', size=12)
plt.show()
fig.savefig('matplotlibExample.png')

# Use Docx to insert the PNG and save report

my_image = doc.add_picture('matplotlibExample.png')  
last_paragraph = doc.paragraphs[-1] 
last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.save("GraphsResults.docx")
